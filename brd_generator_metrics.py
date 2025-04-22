import streamlit as st
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import chromadb
from chromadb.utils import embedding_functions
from chromadb.config import Settings
from uuid import uuid4
import pandas as pd
import plotly.express as px
from streamlit_option_menu import option_menu
import psycopg2
from dotenv import load_dotenv
from nomic import embed
import boto3

# --- Configuration ---
BRD_DOCX_FOLDER = "brd_reports"
os.makedirs(BRD_DOCX_FOLDER, exist_ok=True)

CHROMA_DB_DIR = "./brd_chroma_db"
chroma_client = chromadb.PersistentClient(path=CHROMA_DB_DIR, settings=Settings(
    anonymized_telemetry=False
))
collection = chroma_client.get_or_create_collection(
    name="brds",
    embedding_function=embedding_functions.DefaultEmbeddingFunction()
)

SAP_MODULES = [
    "SD - Sales and Distribution", "MM - Materials Management",
    "FI - Financial Accounting", "CO - Controlling", "PP - Production Planning",
    "WM - Warehouse Management", "PM - Plant Maintenance",
    "QM - Quality Management", "HCM - Human Capital Management", "Other"
]

LLM_MODELS = ["llama3.3:70b", "mistral", "deepseek-r1:32b"]
OLLAMA_API_URL = "http://localhost:11434/api/generate"

VALID_USERS = {
    "admin": "admin123",
    "analyst": "brd@2024",
    "manager": "mgr$789"
}

# --- Login Page ---
def fancy_login():
    st.markdown("""
    <style>
    .stApp {
        background-color: white;
    }
    .login-container {
        max-width: 200px;
        margin: 50px auto;
        padding: 30px;
        background-color: white;
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }
    .login-header {
        text-align: center;
        margin-bottom: 30px;
    }
    .login-title {
        color: #1565c0;
        font-size: 28px;
        font-weight: 600;
        margin-bottom: 5px;
    }
    .login-subtitle {
        color: #7f8c8d;
        font-size: 16px;
    }
    .quick-login {
        background-color: #1e88e5;
        border-radius: 8px;
        padding: 15px;
        margin-top: 20px;
    }
    .forgot-password {
        text-align: center;
        margin-top: 15px;
    }
   .stButton>button {
        background-color: #1e88e5 !important;
        color: white !important;
        border: none !important;
    }
    .stButton>button:hover {
        background-color: #1565c0 !important;
        color: white !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    with st.container():
        st.markdown('<div class="login-container">', unsafe_allow_html=True)
        
        # Header
        st.markdown('<div class="login-header">', unsafe_allow_html=True)
        st.markdown('<div class="login-title">BRD Generator Portal</div>', unsafe_allow_html=True)
        st.markdown('<div class="login-subtitle">Login</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Login Form
        with st.form("login_form"):
            username = st.text_input("**Username**", placeholder="Enter your username")
            password = st.text_input("**Password**", type="password", placeholder="Enter your password")
            submit_button = st.form_submit_button("Login", use_container_width=True)
            
            if submit_button:
                if username in VALID_USERS and password == VALID_USERS[username]:
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.rerun()
                else:
                    st.error("Invalid credentials. Please try again.")
        
        # Quick Login button
        if st.button("Quick Login (Test User)", use_container_width=True):
            st.session_state.logged_in = True
            st.session_state.username = "analyst"
            st.rerun()
        
       # Quick Login Section
       # st.markdown("---")
        #st.markdown("**Test Accounts**")
        
        #test_accounts = pd.DataFrame({
            #"Username": ["admin", "analyst", "manager"],
            #"Password": ["admin123", "brd@2024", "mgr$789"]
        #})
        
        #st.table(test_accounts)#
        
        # Forgot Password
        st.markdown('<div class="forgot-password">', unsafe_allow_html=True)
        st.markdown("[Forgot Password?](#)")
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)

# --- BRD Functions ---
def fetch_existing_brd(ticket_description):
    all_items = collection.get(include=["metadatas", "documents"])
    for i, meta in enumerate(all_items["metadatas"]):
        if meta and "ticket_description" in meta:
            if meta["ticket_description"].strip().lower() == ticket_description.strip().lower():
                return all_items["documents"][i]
    return None

def get_all_brds():
    results = collection.get(include=["metadatas", "documents"])
    return results

def generate_metrics():
    results = get_all_brds()
    if not results or not results["metadatas"]:
        return None
    
    df = pd.DataFrame(results["metadatas"])
    df['ticket_description'] = [meta.get('ticket_description', '') for meta in results["metadatas"]]
    df['timestamp'] = pd.to_datetime(df['timestamp'])
    df['date'] = df['timestamp'].dt.date
    daily_counts = df.groupby('date').size().reset_index(name='count')
    module_dist = df['module'].value_counts().reset_index()
    module_dist.columns = ['Module', 'Count']
    
    return {
        'total_brds': len(df),
        'unique_modules': df['module'].nunique(),
        'daily_counts': daily_counts,
        'module_dist': module_dist,
        'raw_data': df
    }

def generate_brd_content(ticket_description: str, module_tag: str, model_name: str) -> str:
    prompt = f"""
You are a professional SAP business analyst. Create a detailed Business Requirement Document (BRD) for the following ticket.

SAP Module: {module_tag}

Ticket Description (use this as-is for the Problem section):
"{ticket_description}"

Include these sections clearly:
1. Problem: Copy the exact ticket description provided above without any changes.  
2. Analysis  
3. Resolution  
4. Duration (Man Days): Include a table with realistic man days for:
   - Requirement Analysis
   - Functional Development
   - Technical Development
   - Testing
   - Documentation
   - Total
5. Benefit

Return the response in a professional format, structured and ready for documentation.
"""
    response = requests.post(
        OLLAMA_API_URL,
        json={"model": model_name, "prompt": prompt, "stream": False}
    )
    response.raise_for_status()
    return response.json()["response"]

def create_docx(content: str, title: str, module: str) -> str:
    doc = Document()
    doc.add_heading("BUSINESS REQUIREMENT DOCUMENT", 0)
    
    metadata = doc.add_paragraph()
    metadata.add_run("Document Information:\n").bold = True
    metadata.add_run(f"Project: {title}\n")
    metadata.add_run(f"Module: {module}\n")
    metadata.add_run(f"Generated Date: {datetime.now().strftime('%Y-%m-%d')}\n")
    metadata.add_run(f"Generated Time: {datetime.now().strftime('%H:%M:%S')}\n")
    
    doc.add_paragraph()
    
    for line in content.split('\n'):
        if not line.strip():
            continue
        if line.strip().lower().startswith(("problem", "analysis", "resolution", "duration", "benefit")):
            doc.add_heading(line.strip(), level=2)
        elif "|" in line:
            doc.add_paragraph(line.strip(), style='Intense Quote')
        else:
            p = doc.add_paragraph(line.strip())
            p.style.font.size = Pt(11)

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Document generated on {datetime.now().strftime('%Y-%m-%d at %H:%M:%S')}"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    filename = f"BRD_{module.split()[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(BRD_DOCX_FOLDER, filename)
    doc.save(filepath)
    return filepath

def save_to_chroma(brd_text, ticket_description, module, model_name):
    doc_id = str(uuid4())
    metadata = {
        "ticket_description": ticket_description,
        "module": module,
        "model": model_name,
        "timestamp": datetime.now().isoformat()
    }
    collection.add(documents=[brd_text], metadatas=[metadata], ids=[doc_id])

# --- Main Application ---
def main_app():
    st.markdown("""
    <style>
        .main {
            background-color: #f8f9fa;
        }
        .metric-card {
            background-color: white;
            border-radius: 8px;
            padding: 15px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 15px;
        }
        /* Blue button styling */
        .stButton>button {
            background-color: #1e88e5 !important;
            color: white;
            border-radius: 4px;
            border: none;
            padding: 0.5rem 1rem;
            font-weight: 500;
            transition: all 0.3s ease;
        }
        .stButton>button:hover {
            background-color: #1565c0 !important;
            transform: translateY(-1px);
            box-shadow: 0 2px 6px rgba(0,0,0,0.1);
        }
        /* Top navigation styling */
        .top-nav {
            display: flex;
            justify-content: flex-end;
            margin-bottom: 10px;
            align-items: center;
            gap: 20px;
        }
        .app-title {
            position: absolute;
            top: 0.5px;
            left: 10px;
            font-size: 24px;
            font-weight: bold;
            color: #1565c0;
        }
        /* Tabs styling */
        .stTabs [data-baseweb="tab-list"] {
            gap: 10px;
        }
        .stTabs [data-baseweb="tab"] {
            padding: 8px 16px;
            border-radius: 4px;
            transition: all 0.3s ease;
        }
        .stTabs [aria-selected="true"] {
            background-color: #1e88e5 !important;
            color: white !important;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: #e3f2fd !important;
            color: #1e88e5 !important;
        }
    </style>
    """, unsafe_allow_html=True)

    # Application title at top left
    st.markdown('<div class="app-title">BRD Generator Portal</div>', unsafe_allow_html=True)
    
    # Header with navigation at top right
    with st.container():
        st.markdown('<div class="top-nav">', unsafe_allow_html=True)
        
        # Navigation tabs
        selected = option_menu(
            menu_title=None,
            options=["Dashboard", "Create BRD Reports", "BRD Metrics Table"],
            icons=["speedometer", "file-earmark-text", "table"],
            default_index=0,
            orientation="horizontal",
            styles={
                "container": {"padding": "0!important", "background-color": "transparent"},
                "nav-link": {"font-size": "14px", "text-align": "left", "margin": "0px", "--hover-color": "#e3f2fd"},
                "nav-link-selected": {"background-color": "#1e88e5"},
            }
        )
        
        # User info
        st.markdown(f"<div>Welcome, <strong>{st.session_state.username}</strong></div>", 
                   unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    if selected == "Dashboard":
        st.header("📊 Dashboard")
        metrics = generate_metrics()
        
        if not metrics:
            st.warning("No BRDs found in the database. Generate some BRDs first.")
        else:
            col1, col2 = st.columns(2)
            with col1:
                st.markdown('<div class="metric-card"><h3>Total BRDs</h3><h2>{}</h2></div>'.format(metrics['total_brds']), 
                           unsafe_allow_html=True)
            with col2:
                st.markdown('<div class="metric-card"><h3>Unique Modules</h3><h2>{}</h2></div>'.format(metrics['unique_modules']), 
                           unsafe_allow_html=True)
            
            tab1, tab2 = st.tabs(["Daily Activity", "Module Distribution"])
            
            with tab1:
                fig_daily = px.line(metrics['daily_counts'], x='date', y='count',
                                  title="Daily BRD Generation",
                                  labels={'date': 'Date', 'count': 'BRDs Created'},
                                  markers=True)
                st.plotly_chart(fig_daily, use_container_width=True)
            
            with tab2:
                fig_module = px.bar(metrics['module_dist'], x='Module', y='Count',
                                  title="BRDs by SAP Module",
                                  color='Module')
                st.plotly_chart(fig_module, use_container_width=True)
    
    elif selected == "Create BRD Reports":
        st.header("📄 Create New BRD Report")
        
        # Tab for creating new or modifying existing BRDs
        tab1, tab2 = st.tabs(["Generate New BRD", "Upload Modified BRD"])
        
        with tab1:
            col1, col2 = st.columns([3, 1])
            with col1:
                ticket_description = st.text_area("**Ticket Description**", height=200,
                                                placeholder="Enter detailed ticket description...")
            with col2:
                module_choice = st.selectbox("**SAP Module**", SAP_MODULES)
                model_choice = st.selectbox("**LLM Model**", LLM_MODELS)
            
            if st.button("Generate BRD", type="primary", use_container_width=True):
                if not ticket_description.strip():
                    st.warning("Please enter a valid ticket description.")
                else:
                    with st.spinner("Checking for existing BRD..."):
                        existing_brd = fetch_existing_brd(ticket_description)

                    if existing_brd:
                        brd_output = existing_brd
                        st.info("Reusing existing BRD from database")
                    else:
                        with st.spinner("Generating BRD..."):
                            brd_output = generate_brd_content(ticket_description, module_choice, model_choice)
                            save_to_chroma(brd_output, ticket_description, module_choice, model_choice)
                            st.success("BRD Generated Successfully!")

                    title = f"BRD - {module_choice.split()[0]}"
                    
                    with st.expander("View BRD Content"):
                        st.code(brd_output, language='markdown')

                    file_path = create_docx(brd_output, title, module_choice)
                    with open(file_path, "rb") as f:
                        st.download_button("Download BRD (.docx)", f, 
                                         file_name=os.path.basename(file_path),
                                         use_container_width=True)
        
        with tab2:
            st.markdown("### Upload Modified BRD Report")
            uploaded_file = st.file_uploader("Choose a BRD file (.docx)", type="docx")
            
            if uploaded_file is not None:
                try:
                    doc = Document(uploaded_file)
                    brd_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    
                    # Extract ticket description (assuming it's in the Problem section)
                    problem_section = ""
                    for i, para in enumerate(doc.paragraphs):
                        if para.text.strip().lower().startswith("problem"):
                            problem_section = doc.paragraphs[i+1].text
                            break
                    
                    if problem_section:
                        st.markdown("**Found Ticket Description:**")
                        st.info(problem_section)
                        
                        if st.button("Update BRD in Database", type="primary"):
                            with st.spinner("Updating BRD..."):
                                if update_brd_in_chroma(brd_content, problem_section):
                                    st.success("BRD updated successfully in database!")
                                else:
                                    st.warning("No matching BRD found to update. Creating new entry...")
                                    module = "Other"  # Default module if we can't determine from upload
                                    for para in doc.paragraphs:
                                        if "Module:" in para.text:
                                            module = para.text.split("Module:")[1].strip()
                                            break
                                    save_to_chroma(brd_content, problem_section, module, "manual_upload")
                                    st.success("New BRD created in database!")
                    else:
                        st.error("Could not identify Problem section in the uploaded document.")
                        
                    with st.expander("View Uploaded BRD Content"):
                        st.code(brd_content, language='markdown')
                        
                except Exception as e:
                    st.error(f"Error processing file: {str(e)}")
    
    elif selected == "BRD Metrics Table":
        st.header("📋 BRD Metrics Table")
        metrics = generate_metrics()
        
        if not metrics:
            st.warning("No BRDs found in the database. Generate some BRDs first.")
        else:
            st.dataframe(metrics['raw_data'][['module', 'model', 'ticket_description', 'timestamp']]
                         .sort_values('timestamp', ascending=False),
                         use_container_width=True,
                         column_config={
                             "timestamp": st.column_config.DatetimeColumn("Timestamp"),
                             "ticket_description": st.column_config.TextColumn("Ticket Description", width="large")
                         })
            
            csv = metrics['raw_data'].to_csv(index=False).encode('utf-8')
            st.download_button("Export Data as CSV", csv,
                             file_name="brd_metrics.csv",
                             mime="text/csv", 
                             use_container_width=True)

# --- App Execution ---
st.set_page_config(page_title="SAP BRD Generator", layout="wide", page_icon="📊")

if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

if st.session_state.logged_in:
    main_app()
else:
    fancy_login()
